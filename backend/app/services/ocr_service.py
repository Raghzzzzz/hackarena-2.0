import email
import io
import mimetypes
import re
import uuid
from pathlib import Path
from typing import Any

import fitz
import numpy as np
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image

from ..database import UPLOAD_DIR

_ocr_engine = None


def get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        from paddleocr import PaddleOCR

        _ocr_engine = PaddleOCR(
            use_angle_cls=True,
            lang="en",
            show_log=False,
            use_gpu=False,
        )
    return _ocr_engine


def save_upload(file_bytes: bytes, filename: str) -> tuple[str, str, int]:
    ext = Path(filename).suffix.lower()
    stored_name = f"{uuid.uuid4().hex}{ext}"
    dest = UPLOAD_DIR / stored_name
    dest.write_bytes(file_bytes)
    mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    return str(dest), mime_type, len(file_bytes)


def _image_to_array(image: Image.Image) -> np.ndarray:
    if image.mode != "RGB":
        image = image.convert("RGB")
    return np.array(image)


def _parse_paddle_result(result) -> tuple[str, list[dict]]:
    lines: list[str] = []
    entities: list[dict] = []

    if not result:
        return "", entities

    for page in result:
        if not page:
            continue
        for line in page:
            bbox_points, (text, conf) = line
            text = (text or "").strip()
            if not text:
                continue
            lines.append(text)
            xs = [p[0] for p in bbox_points]
            ys = [p[1] for p in bbox_points]
            bbox = [float(min(xs)), float(min(ys)), float(max(xs)), float(max(ys))]
            from .field_extractor import classify_entities

            entities.append(
                {
                    "type": classify_entities(text, bbox, conf),
                    "text": text,
                    "conf": round(float(conf), 4),
                    "bbox": bbox,
                }
            )

    return "\n".join(lines), entities


def _run_paddle_on_image(image: Image.Image) -> tuple[str, list[dict]]:
    ocr = get_ocr_engine()
    arr = _image_to_array(image)
    try:
        result = ocr.ocr(arr, cls=True)
        return _parse_paddle_result(result)
    except Exception:
        return "", []


def _extract_pdf_text(path: Path) -> str:
    doc = fitz.open(path)
    try:
        parts = [page.get_text().strip() for page in doc if page.get_text().strip()]
        return "\n".join(parts)
    finally:
        doc.close()


def _ocr_images(images: list[Image.Image]) -> tuple[str, list[dict]]:
    all_text: list[str] = []
    all_entities: list[dict] = []
    y_offset = 0

    for image in images:
        page_text, page_entities = _run_paddle_on_image(image)
        if page_text:
            all_text.append(page_text)
        for entity in page_entities:
            bbox = entity["bbox"]
            entity = {**entity, "bbox": [bbox[0], bbox[1] + y_offset, bbox[2], bbox[3] + y_offset]}
            all_entities.append(entity)
        y_offset += image.height

    return "\n\n".join(all_text), all_entities


def _pdf_to_images(path: Path, dpi: int = 200) -> list[Image.Image]:
    images: list[Image.Image] = []
    doc = fitz.open(path)
    try:
        for page in doc:
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            images.append(Image.frombytes("RGB", [pix.width, pix.height], pix.samples))
    finally:
        doc.close()
    return images


def _read_plain_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    doc = DocxDocument(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_xlsx(path: Path) -> str:
    wb = load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = []
    try:
        for sheet in wb.worksheets:
            lines.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(lines)


def _entities_from_text(text: str) -> list[dict]:
    entities: list[dict] = []
    y = 20
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        from .field_extractor import classify_entities

        entities.append(
            {
                "type": classify_entities(line, [0, y, 400, y + 16], 0.99),
                "text": line,
                "conf": 0.99,
                "bbox": [20, y, 400, y + 16],
            }
        )
        y += 20
    return entities


def _process_eml(file_bytes: bytes) -> tuple[str, list[dict], list[tuple[str, bytes]]]:
    msg = email.message_from_bytes(file_bytes)
    body_parts: list[str] = []
    attachments: list[tuple[str, bytes]] = []

    if msg.get("Subject"):
        body_parts.append(f"Subject: {msg['Subject']}")
    if msg.get("From"):
        body_parts.append(f"From: {msg['From']}")

    if msg.is_multipart():
        for part in msg.walk():
            content_disposition = part.get("Content-Disposition", "")
            filename = part.get_filename()
            payload = part.get_payload(decode=True)
            if filename and payload:
                attachments.append((filename, payload))
            elif part.get_content_type() == "text/plain" and payload:
                body_parts.append(payload.decode(part.get_content_charset() or "utf-8", errors="replace"))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            body_parts.append(payload.decode(msg.get_content_charset() or "utf-8", errors="replace"))

    body_text = "\n".join(body_parts)
    entities = _entities_from_text(body_text)
    return body_text, entities, attachments


def process_document(storage_path: str, mime_type: str, filename: str) -> dict[str, Any]:
    path = Path(storage_path)
    ext = path.suffix.lower()
    filename_lower = filename.lower()

    # Plain text / notes
    if ext in {".txt", ".md", ".csv", ".json"} or mime_type.startswith("text/"):
        text = _read_plain_text(path)
        entities = _entities_from_text(text)
        from .field_extractor import build_ocr_result, extract_fields_from_text

        ocr_result = build_ocr_result(text, entities)
        fields = extract_fields_from_text(text, entities)
        return {"raw_text": text, "ocr_result": ocr_result, "extracted_fields": fields}

    # Word documents
    if ext in {".docx"}:
        text = _read_docx(path)
        entities = _entities_from_text(text)
        from .field_extractor import build_ocr_result, extract_fields_from_text

        ocr_result = build_ocr_result(text, entities)
        fields = extract_fields_from_text(text, entities)
        return {"raw_text": text, "ocr_result": ocr_result, "extracted_fields": fields}

    # Excel spreadsheets
    if ext in {".xlsx", ".xlsm"} or "spreadsheet" in mime_type:
        text = _read_xlsx(path)
        entities = _entities_from_text(text)
        from .field_extractor import build_ocr_result, extract_fields_from_text

        ocr_result = build_ocr_result(text, entities)
        fields = extract_fields_from_text(text, entities)
        return {"raw_text": text, "ocr_result": ocr_result, "extracted_fields": fields}

    # Email with optional attachment OCR
    if ext == ".eml" or mime_type == "message/rfc822":
        file_bytes = path.read_bytes()
        body_text, body_entities, attachments = _process_eml(file_bytes)
        combined_text = body_text
        combined_entities = list(body_entities)
        y_offset = len(body_entities) * 20 + 40

        for att_name, att_bytes in attachments:
            att_ext = Path(att_name).suffix.lower()
            temp_path = UPLOAD_DIR / f"eml_att_{uuid.uuid4().hex}{att_ext}"
            temp_path.write_bytes(att_bytes)
            try:
                att_result = process_document(str(temp_path), mimetypes.guess_type(att_name)[0] or "", att_name)
                if att_result.get("raw_text"):
                    combined_text += f"\n\n--- Attachment: {att_name} ---\n{att_result['raw_text']}"
                for entity in att_result.get("ocr_result", {}).get("entities", []):
                    bbox = entity["bbox"]
                    entity = {**entity, "bbox": [bbox[0], bbox[1] + y_offset, bbox[2], bbox[3] + y_offset]}
                    combined_entities.append(entity)
                y_offset += 400
            finally:
                temp_path.unlink(missing_ok=True)

        from .field_extractor import build_ocr_result, extract_fields_from_text

        ocr_result = build_ocr_result(combined_text, combined_entities)
        fields = extract_fields_from_text(combined_text, combined_entities)
        return {"raw_text": combined_text, "ocr_result": ocr_result, "extracted_fields": fields}

    # PDF — try text layer first, then OCR rendered pages
    if ext == ".pdf" or mime_type == "application/pdf":
        text_layer = _extract_pdf_text(path)
        if text_layer and len(text_layer.strip()) > 20:
            entities = _entities_from_text(text_layer)
            from .field_extractor import build_ocr_result, extract_fields_from_text

            ocr_result = build_ocr_result(text_layer, entities)
            fields = extract_fields_from_text(text_layer, entities)
            return {"raw_text": text_layer, "ocr_result": ocr_result, "extracted_fields": fields}

        images = _pdf_to_images(path)
        text, entities = _ocr_images(images)
        if not text.strip() and text_layer:
            text = text_layer
            entities = _entities_from_text(text_layer)
        from .field_extractor import build_ocr_result, extract_fields_from_text

        ocr_result = build_ocr_result(text, entities)
        fields = extract_fields_from_text(text, entities)
        return {"raw_text": text, "ocr_result": ocr_result, "extracted_fields": fields}

    # Images — including handwritten scans
    if ext in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"} or mime_type.startswith("image/"):
        image = Image.open(path)
        text, entities = _run_paddle_on_image(image)
        from .field_extractor import build_ocr_result, extract_fields_from_text

        ocr_result = build_ocr_result(text, entities)
        fields = extract_fields_from_text(text, entities)
        return {"raw_text": text, "ocr_result": ocr_result, "extracted_fields": fields}

    raise ValueError(f"Unsupported file format: {filename} ({mime_type})")
